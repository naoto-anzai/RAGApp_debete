from pathlib import Path
from docx import Document

PROJECT_DIR = Path(__file__).resolve().parents[1]

from e10_report_agent.report_state import ReportState

def create_save_word_node():
    def save_word_node(state: ReportState):

        title = state["title"]
        report = state["report"]
        file_name = "report.docx"

        safe_file_name = Path(file_name).name

        if not safe_file_name.lower().endswith(".docx"):
            safe_file_name += ".docx"

        output_path = PROJECT_DIR / safe_file_name

        document = Document()

        document.add_heading(title, level=0)

        for line in report.splitlines():
            stripped_line = line.strip()

            if not stripped_line:
                continue

            if stripped_line.startswith("## "):
                document.add_heading(
                    stripped_line.removeprefix("## "),
                    level=1
                )
            elif stripped_line.startswith("### "):
                document.add_heading(
                    stripped_line.removeprefix("### "),
                    level=2
                )
            else:
                document.add_paragraph(stripped_line)

        document.save(output_path)

        print(f"Wordファイルを保存しました: {output_path}")
    return save_word_node